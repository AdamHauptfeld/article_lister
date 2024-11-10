import pandas as pd
import docx
import PyPDF2
from langchain_ollama import OllamaLLM
from langchain_core.prompts import ChatPromptTemplate


template = (
    "You are tasked with extracting keywords from the following text content: {file_content}. "
    "Please analyze this article and return exactly 5 keywords or short phrases (of no more than 5 words) that best represent its main themes and concepts. \
    Format your response as a simple comma-separated list with no additional explanation or commentary."
)

model = OllamaLLM(model="llama3.1")


def keyworder_main(file_path):
    file_content = file_categorizor(file_path)
    keywords = parse_with_ollama(file_content)
    return(keywords)


#determines the file type and extracts the file's content
def file_categorizor(file_path):
    file_extension = file_path.rsplit(".")[-1].lower()

    if file_extension == "docx":
        document = docx.Document(file_path)
        file_content = '\n'.join([paragraph.text for paragraph in document.paragraphs])
        return file_content
    elif file_extension == "doc":
        document = docx.Document(file_path)
        file_content = '\n'.join([paragraph.text for paragraph in document.paragraphs])
        return file_content
    elif file_extension == "txt":
        file_content = file_path.read()
        return file_content
    elif file_extension == "pdf":
        document = PyPDF2.PdfReader(file_path)
        file_content = ''
        for page in document.pages:
            file_content += page.extract_text()
        return file_content
    elif file_extension == "xlsx":
        df = pd.read_excel(file_path)
        file_content = df.to_string()
        return file_content
    else:
        print("Error: unacceptable file type.")


#uses llama3.1 to find 5 keywords and phrases that summarizes the file's content.
def parse_with_ollama(file_content):
    if file_content == '' or None:
        response = None
    
    else:       
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | model
        response = chain.invoke(file_content)

    return response